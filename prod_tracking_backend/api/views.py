from django.shortcuts import render, get_object_or_404
from rest_framework.decorators import api_view
from django.core.files.storage import FileSystemStorage
# Create your views here.
from django.http import HttpResponse
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .models import Product
from django.http import Http404
from . serializers import ProductSerializers
import docx
from docx.shared import Inches
from io import StringIO
from io import BytesIO

class ProductList(APIView):
    def get(self, request):
        product = Product.objects.all()
        serializer = ProductSerializers(product, many=True)
        return Response(serializer.data)

    def post(self, request, format=None):
        serializer = ProductSerializers(data = request.data)

        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        print(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class ProductDetail(APIView):
    # def get_object(self, description, datetime):
    #     try:
    #         return Product.objects.get(datetime=datetime, description=description)
    #     except Product.DoesNotExist:
    #         raise Http404

    def get_object(self, id):
        try:
            return Product.objects.get(id=id)
        except Product.DoesNotExist:
            raise Http404

    # def get(self, request,desc, datetime, format=None):
    #     product = self.get_object(desc, datetime)
    #     product = ProductSerializers(product)
    #     return Response(product.data)

    def get(self, request,id=None, desc=None, datetime=None, format=None):
        if desc == None and datetime == None:
            product = self.get_object(id)
            product = ProductSerializers(product)
            return Response(product.data)
        else:
            product = Product.objects.filter(datetime=datetime, description=desc)
            serializer = ProductSerializers(product, many=True)
            return Response(serializer.data)

    def put(self, request,id, format=None):
        product = self.get_object(id)
        serializer = ProductSerializers(product, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    # def delete(self, request, desc, datetime, format=None):
    #     product = self.get_object(desc, datetime)
    #     product.delete()
    #     return Response(status=status.HTTP_204_NO_CONTENT)

    def delete(self, request, id, format=None):
        product = self.get_object(id)
        product.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

# function base views start here
@api_view(['GET', 'POST'])
def load_from_file(request):
    uploaded_file_url = ""
    if request.method == 'POST' and request.FILES['uploadedfile']:
        myfile = request.FILES['uploadedfile']
        fs = FileSystemStorage()
        filename = fs.save(myfile.name, myfile)
        uploaded_file_url = fs.url(filename)
        print(uploaded_file_url)

    # filename = "./static/input.txt.docx"
    doc = docx.Document("."+uploaded_file_url)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    for text in fullText[1:]:
        data = text.split("\t")
        print(data)
        prod = Product.objects.create(description=data[1],datetime=data[2],longitude=data[3],latitude=data[4],elevation=data[5])

    print(data)
    # with open(myfile, 'r') as f:
    #     content = f.readlines()
    #     for line in content[1:]:
    #         data = line.split()
    #     print(data)

    return HttpResponse("ok")

@api_view(['GET', 'POST'])
def download_docx(request):
    document = docx.Document()
    products = Product.objects.all()

    fields = [field.name for field in Product._meta.fields]

    document.add_paragraph("{}\t{}\t{}\t{}\t{}\t{}".format(*fields))
    for prod in products:
        document.add_paragraph("{}\t{}\t{}\t{}\t{}\t{}\n".format(
            prod.id,
            prod.description,
            prod.datetime,
            prod.longitude,
            prod.latitude,
            prod.elevation
        ))

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=download.docx'
    document.save(response)

    return response
